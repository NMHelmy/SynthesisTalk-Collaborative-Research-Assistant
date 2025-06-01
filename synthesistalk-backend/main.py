from fastapi import FastAPI, Request, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from collections import defaultdict
from duckduckgo_search import DDGS
from routers import files, upload, extract
from dotenv import load_dotenv
import os
import requests
import re
from datetime import datetime
from typing import List, Dict
from docx import Document
from reportlab.pdfgen import canvas
from routers import upload, files, extract
from fastapi.responses import JSONResponse

# Load environment variables
load_dotenv()
API_KEY = os.getenv("NGU_API_KEY")
BASE_URL = os.getenv("NGU_BASE_URL")
MODEL = os.getenv("NGU_MODEL")

# Initialize FastAPI app
app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files for uploaded content
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

# Register routers
app.include_router(files.router, prefix="/api")
app.include_router(upload.router, prefix="/api")
app.include_router(extract.router, prefix="/api")

@app.get("/")
def root():
    return {"status": "SynthesisTalk Backend Running"}

from state import session_histories, session_summaries, session_extracted

def load_chat_history(session_id):
    return session_histories.get(session_id, [])

def load_summaries(session_id):
    return session_summaries.get(session_id, [])

def load_extracted_text(session_id):
    return session_extracted.get(session_id, "")

def generate_pdf(path, chat_data, summaries, extracted_text):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4

    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    y = height - 50
    line_height = 15

    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, y, "SynthesisTalk Report")
    y -= 40

    # Conversation section
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Conversation")
    y -= 25

    c.setFont("Helvetica", 12)
    for item in chat_data:
        if item.get("type") == "chart":
            continue  # Skip visualized data
        text = f"{item['role'].capitalize()}: {item['content']}"
        for line in text.split("\n"):
            if y <= 40:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 12)
            c.drawString(50, y, line)
            y -= line_height

    c.save()

def generate_docx(path, chat_data, summaries, extracted_text):
    doc = Document()
    doc.add_heading("SynthesisTalk Report", 0)

    doc.add_heading("Conversation", level=1)
    for item in chat_data:
        doc.add_paragraph(f"{item['role'].capitalize()}: {item['content']}")

    doc.save(path)

# Input schema
class ChatRequest(BaseModel):
    session_id: str
    prompt: str
    mode: str = "normal"  # normal, cot, react

class SearchPayload(BaseModel):
    query: str

# System prompt logic

# Web search endpoint (for Web Search button)
class SearchPayload(BaseModel):
    query: str

# Mode-specific system prompts
def get_system_prompt(mode: str) -> str:
    if mode == 'normal':
        return "You are a helpful assistant. Provide concise, direct answers without revealing your reasoning."
    elif mode == 'cot':
        return (
            "Let's think step by step. Provide a clear reasoning chain to solve the problem."
            " Use numbers, equations, and logic when needed. Finish with 'The answer is ...'."
        )
        
    elif mode == 'react':
        return (
            "You are an AI assistant that uses ReAct (Reasoning + Acting).\n"
            "When you face a question requiring up-to-date or external info, follow this format:\n\n"
            "Thought: [reasoning]\n"
            "Action: search[query]\n\n"
            "After receiving an observation from the tool, continue with:\n"
            "Observation: [real result]\n"
            "Thought: [reflect on the observation]\n"
            "Answer: [final answer based ONLY on the observation]\n\n"
            "Do not say you lack access to real-time data if a search tool is available.\n"
            "Always complete the Thought → Action → Observation → Answer flow."
            "Always wait for the observation before answering."
            "DO NOT answer before receiving an observation.\n"
        )
    return "You are a helpful assistant."

# DuckDuckGo search tool
def search_web(query: str) -> str:
    with DDGS() as ddgs:
        results = ddgs.text(query, max_results=5)
    lines = []
    for r in results:
        title = r.get("title", "(no title)")
        body = r.get("body", "").strip()
        href = r.get("href", "")
        lines.append(f"- <a href='{href}' target='_blank'><strong>{title}</strong></a><br>{body}")

    return "<br><br>".join(lines) or "No results found."

# Core LLM call wrapper
def call_llm(messages: list) -> str:
    payload = {
        "model": MODEL,
        "messages": messages,
        "temperature": 0.7,
        "top_p": 0.9
    }
    resp = requests.post(f"{BASE_URL}/chat/completions", json=payload, headers= {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    })
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]

class RestoreRequest(BaseModel):
    session_id: str
    messages: List[Dict[str, str]]

@app.post("/api/restore")
async def restore_chat_history(req: RestoreRequest):
    session_histories[req.session_id] = req.messages
    return {"status": "restored", "count": len(req.messages)}

# Chat endpoint supporting Normal, CoT, and ReAct
@app.post("/api/chat")
async def chat(req: ChatRequest):
    sid = req.session_id
    mode = req.mode.lower()
    prompt = req.prompt.strip()

    # # Handle direct date queries without LLM
    # if re.search(r"\b(?:date|today)\b", prompt.lower()):
    #     today = datetime.now().strftime("%B %d, %Y")
    #     reply = f"Today's date is {today}."
    #     # Persist reply
    #     session_histories[sid].append({"role": "assistant", "content": reply})
    #     return {"response": reply}


    # Retrieve persisted history (user & assistant only)
    history = session_histories[sid]
    system_msg = {"role": "system", "content": get_system_prompt(mode)}
    messages = [system_msg] + history + [{"role": "user", "content": prompt}]

    if mode == 'react':
        initial_response = call_llm(messages)
        action_match = re.search(r"Action:\s*(\w+)\[(.*?)\]", initial_response)

        if action_match and action_match.group(1).lower() == 'search':
            query = action_match.group(2).strip()
            observation = search_web(query)

            # Format a cleaner Observation for the LLM
            observation_msg = f"Observation: {observation}"

            # Append LLM's Action response + Observation
            messages.append({"role": "assistant", "content": initial_response})
            messages.append({"role": "assistant", "content": observation_msg})

            # Second LLM call: reflect and respond based on the real observation
            followup_response = call_llm(messages)

            reply = f"{initial_response}\n\n{observation_msg}\n\n{followup_response}"

        else:
            reply = initial_response

    elif mode == 'cot':
        response = call_llm(messages)
        # Block any ReAct-style behavior in CoT mode
        if "Action:" in response or "Observation:" in response:
            # Remove any such patterns if they sneak in
            response = re.sub(r"Action:.*?\n", "", response)
            response = re.sub(r"Observation:.*?\n", "", response)

        # Add CoT prefix only if not already present
        reply = response
        if not reply.lower().startswith("let's think") and "step" not in reply.lower():
            reply = "Let's think step by step.\n" + reply



    elif mode == 'normal':
        reply = call_llm(messages)

        # Save summary if prompt includes summarize
        if "summarize" in prompt.lower():
            points = [p.strip("•- ") for p in reply.split("\n") if p.strip()]
            session_summaries[sid] = points


    history.append({"role": "user", "content": prompt})
    history.append({"role": "assistant", "content": reply})

    return {"response": reply}

@app.post("/api/search")
async def search(payload: SearchPayload):
    result = search_web(payload.query)
    return {"results": result}

@app.post("/api/visualize")
async def visualize(req: Dict):
    text = req.get("text", "").strip()
    if not text:
        return {"data": []}
    # Try to extract lines like "Country: Number unit"
    pattern = r"(?:\d+\.\s*)?([\w\s]+):\s*([\d.,]+)\s*(trillion|billion|million)?"
    matches = re.findall(pattern, text, flags=re.IGNORECASE)

    insights = []
    for name, num, unit in matches:
        try:
            value = float(num.replace(",", ""))
            if unit:
                unit = unit.lower()
                if unit == "trillion":
                    value *= 1_000_000_000_000
                elif unit == "billion":
                    value *= 1_000_000_000
                elif unit == "million":
                    value *= 1_000_000
            insights.append({"label": name.strip(), "value": int(value)})
        except:
            continue
    return {"data": insights[:5]}  # limit to 5 results

@app.post("/api/export")
async def export_report(request: Request):
    data = await request.json()
    session_id = data.get("session_id")
    export_format = data.get("format", "pdf")

    chat_data = load_chat_history(session_id)
    summaries = load_summaries(session_id)
    extracted_text = load_extracted_text(session_id)

    if not chat_data:
        return JSONResponse({"error": "No chat data found for this session."}, status_code=400)

    # Use first user message as filename
    first_user_message = next((item['content'] for item in chat_data if item['role'] == 'user'), "SynthesisTalk")
    safe_title = re.sub(r'[^a-zA-Z0-9_]+', '_', first_user_message.strip())[:40]
    filename_base = safe_title or f"SynthesisTalk_{session_id}"
    filename = f"{filename_base}.{export_format}"
    filepath = os.path.join("exports", filename)
    os.makedirs("exports", exist_ok=True)

    try:
        if export_format == "pdf":
            generate_pdf(filepath, chat_data, summaries, extracted_text)
        elif export_format == "docx":
            generate_docx(filepath, chat_data, summaries, extracted_text)
        else:
            return JSONResponse({"error": "Unsupported format"}, status_code=400)

        if not os.path.exists(filepath) or os.path.getsize(filepath) == 0:
            return JSONResponse({"error": "Generated file is empty or missing"}, status_code=500)

        return FileResponse(path=filepath, filename=filename, media_type="application/octet-stream")
    except Exception as e:
        return JSONResponse({"error": f"Export failed: {str(e)}"}, status_code=500)
